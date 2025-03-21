import easyocr
import fitz  # PyMuPDF
import docx
import os
from pathlib import Path
from tqdm import tqdm
import time

def pdf_to_docx_with_easyocr(pdf_path, output_docx=None, temp_dir="temp_images"):
    """
    Convertit un PDF non sélectionnable en document Word en utilisant EasyOCR
    
    Args:
        pdf_path: Chemin vers le fichier PDF
        output_docx: Chemin de sortie pour le fichier Word (optionnel)
        temp_dir: Dossier pour les images temporaires
    """
    start_time = time.time()
    
    # Créer le dossier temporaire s'il n'existe pas
    Path(temp_dir).mkdir(exist_ok=True)
    
    if output_docx is None:
        output_docx = pdf_path.replace('.pdf', '.docx')
    
    print(f"Initialisation d'EasyOCR (cela peut prendre un moment la première fois)...")
    # Initialiser le lecteur OCR avec le français
    reader = easyocr.Reader(['fr'])
    
    # Ouvrir le PDF
    print(f"Ouverture du PDF: {pdf_path}")
    pdf_document = fitz.open(pdf_path)
    doc = docx.Document()
    
    # Ajouter un titre au document Word
    pdf_name = os.path.basename(pdf_path).replace('.pdf', '')
    doc.add_heading(pdf_name, 0)
    
    # Traiter chaque page
    total_pages = len(pdf_document)
    print(f"Traitement de {total_pages} pages...")
    
    for page_num in tqdm(range(total_pages), desc="Pages traitées"):
        page = pdf_document.load_page(page_num)
        
        # Ajouter un en-tête de page
        doc.add_heading(f"Page {page_num+1}", level=2)
        
        # Extraire l'image de la page avec une haute résolution pour meilleur OCR
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        image_path = os.path.join(temp_dir, f"page_{page_num}.png")
        pix.save(image_path)
        
        # Appliquer l'OCR
        result = reader.readtext(image_path)
        
        # Extraire le texte et l'ajouter au document Word
        current_paragraph = None
        
        for detection in result:
            text = detection[1]
            confidence = detection[2]
            
            # Ajouter le texte au document
            if len(text) > 50:  # Paragraphe plus long
                p = doc.add_paragraph(text)
                current_paragraph = None
            elif current_paragraph is None:  # Nouveau texte court
                current_paragraph = doc.add_paragraph(text)
            else:  # Continuer un paragraphe court existant
                current_paragraph.add_run(" " + text)
        
        # Ajouter un saut de page après chaque page sauf la dernière
        if page_num < total_pages - 1:
            doc.add_page_break()
    
    # Enregistrer le document Word
    doc.save(output_docx)
    
    # Nettoyer les fichiers temporaires
    for page_num in range(total_pages):
        image_path = os.path.join(temp_dir, f"page_{page_num}.png")
        if os.path.exists(image_path):
            os.remove(image_path)
    
    elapsed_time = time.time() - start_time
    print(f"Conversion terminée en {elapsed_time:.2f} secondes.")
    print(f"Document Word créé: {output_docx}")
    return output_docx

if __name__ == "__main__":
    # Liste des PDFs à traiter
    pdf_files = [
        "AUPSRVE-2023_fr.pdf",
        # Ajoutez d'autres fichiers ici
    ]
    
    for pdf_file in pdf_files:
        pdf_to_docx_with_easyocr(pdf_file)