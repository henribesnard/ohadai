import os
import re
import json
import glob
from typing import List, Dict, Any, Optional, Set, Tuple
from tqdm import tqdm
import argparse
from pathlib import Path
from docx import Document

# Import des classes définies précédemment
from src.vector_db.ohada_vector_db_structure import OhadaVectorDB, OhadaDocument, OhadaMetadataBuilder, OhadaReference, OhadaEmbedder

class OhadaWordProcessor:
    """Processeur de Word pour le plan comptable OHADA"""
    
    def __init__(self, vector_db: OhadaVectorDB, embedder = None):
        """Initialise le processeur de documents Word
        
        Args:
            vector_db: Instance de la base de données vectorielle
            embedder: Processeur d'embeddings à utiliser
        """
        self.vector_db = vector_db
        
        # Utiliser l'embedder de la base vectorielle si aucun n'est fourni
        if embedder is None:
            self.embedder = self.vector_db.embedder
        else:
            self.embedder = embedder
        
        # Charger la table des matières
        self.toc = self._load_toc()
    
    def _load_toc(self) -> Dict[str, Any]:
        """Charge la table des matières structurée"""
        return self.vector_db.toc
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extrait le texte d'un fichier Word
        
        Args:
            docx_path: Chemin vers le fichier Word
            
        Returns:
            Texte extrait du document Word
        """
        try:
            # Ouvrir le document Word
            doc = Document(docx_path)
            text_parts = []
            
            # Extraire le texte de chaque paragraphe
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(row_text))
            
            # Joindre toutes les parties de texte
            full_text = "\n".join(text_parts)
            
            # Nettoyer le texte (supprimer les caractères spéciaux problématiques)
            full_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', full_text)
            
            return full_text
        except Exception as e:
            print(f"Erreur lors de l'extraction du texte du document Word {docx_path}: {e}")
            return ""
    
    def build_metadata_from_docx_path(self, docx_path: str) -> Dict[str, Any]:
        """Construit les métadonnées à partir du chemin d'un fichier Word
        
        Args:
            docx_path: Chemin vers le fichier Word
            
        Returns:
            Métadonnées structurées
        """
        metadata = {}
        
        # Extraire les informations du nom de fichier
        file_name = os.path.basename(docx_path)
        dir_name = os.path.dirname(docx_path)
        
        # Déterminer si c'est un document de présentation OHADA
        if "presentation_ohada" in dir_name:
            metadata["document_type"] = "presentation_ohada"
            # Générer un ID basé sur le nom du fichier
            file_base = os.path.splitext(file_name)[0].lower()
            # Nettoyer le nom pour l'ID
            clean_id = re.sub(r'[^a-z0-9_]', '_', file_base)
            metadata["id"] = f"presentation_ohada_{clean_id}"
            # Utiliser le nom du fichier comme titre
            metadata["title"] = file_name.replace('.docx', '')
            return metadata
        
        # Pour les documents standard du plan comptable
        # Déterminer la partie
        partie_match = re.search(r'partie_(\d+)', dir_name)
        if partie_match:
            metadata["partie"] = int(partie_match.group(1))
        
        # Déterminer le chapitre
        chapitre_match = re.search(r'chapitre_(\d+)', file_name)
        if chapitre_match:
            metadata["chapitre"] = int(chapitre_match.group(1))
        
        # Générer un ID
        if "partie" in metadata and "chapitre" in metadata:
            metadata["id"] = f"partie_{metadata['partie']}_chapitre_{metadata['chapitre']}"
            metadata["document_type"] = "chapitre"
            metadata["parent_id"] = f"partie_{metadata['partie']}"
        
        # Rechercher des informations supplémentaires dans la table des matières si disponible
        if self.toc and "lookup" in self.toc and metadata.get("id") in self.toc["lookup"]:
            toc_entry = self.toc["lookup"][metadata["id"]]
            metadata["title"] = toc_entry.get("title", f"Chapitre {metadata.get('chapitre', '')}")
            metadata["page_debut"] = toc_entry.get("start_page", toc_entry.get("page"))
            metadata["page_fin"] = toc_entry.get("end_page")
        else:
            # Essayer d'extraire le titre du document Word
            try:
                doc = Document(docx_path)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        title_match = re.search(r'(?:CHAPITRE|Chapitre)\s+\d+\s*[:\-]?\s*(.+)', paragraph.text)
                        if title_match:
                            title = title_match.group(1).strip()
                            metadata["title"] = f"Chapitre {metadata.get('chapitre', '')}: {title}"
                            break
            except Exception as e:
                print(f"Erreur lors de l'extraction du titre du document {docx_path}: {e}")
            
            # Titre par défaut si non trouvé
            if "title" not in metadata:
                metadata["title"] = f"Chapitre {metadata.get('chapitre', '')}"
        
        # Ajouter le chemin du fichier Word
        metadata["docx_path"] = docx_path
        
        return metadata
    
    def process_docx_file(self, docx_path: str) -> Optional[OhadaDocument]:
        """Traite un fichier Word et crée un document OHADA
        
        Args:
            docx_path: Chemin vers le fichier Word
            
        Returns:
            Document OHADA ou None en cas d'erreur
        """
        print(f"Traitement du fichier Word: {docx_path}")
        
        # Extraire le texte du document Word
        text = self.extract_text_from_docx(docx_path)
        
        if not text or len(text.strip()) < 100:  # Vérifier qu'il y a du contenu significatif
            print(f"Avertissement: Texte insuffisant extrait de {docx_path}")
            return None
        
        # Construire les métadonnées à partir du chemin du fichier
        metadata = self.build_metadata_from_docx_path(docx_path)
        
        # Générer un ID
        doc_id = metadata.get("id", os.path.basename(docx_path).replace('.docx', ''))
        
        # Déterminer le type de document (chapitre standard ou présentation OHADA)
        if metadata.get("document_type") == "presentation_ohada":
            # Pour les documents de présentation OHADA
            reference = OhadaReference(
                partie_num=0,  # Pas de partie pour les documents de présentation
                chapitre_num=None,
                chapitre_title=metadata.get("title")
            )
        else:
            # Pour les chapitres standard
            reference = OhadaReference(
                partie_num=metadata.get("partie", 1),
                chapitre_num=metadata.get("chapitre"),
                chapitre_title=metadata.get("title")
            )
        
        # Créer le document
        document = OhadaDocument(
            id=doc_id,
            text=text,
            metadata=metadata,
            reference=reference,
            pdf_path=None  # On ne stocke pas de chemin PDF ici
        )
        
        return document
    
    def ingest_docx_file(self, docx_path: str) -> bool:
        """Ingère un fichier Word dans la base vectorielle
        
        Args:
            docx_path: Chemin vers le fichier Word
            
        Returns:
            True si l'ingestion a réussi, False sinon
        """
        # Traiter le fichier
        document = self.process_docx_file(docx_path)
        
        if not document:
            return False
        
        # Générer l'embedding
        print(f"Génération de l'embedding pour {os.path.basename(docx_path)}...")
        embedding = self.embedder.generate_embedding(document.text)
        
        # Déterminer les collections pour ce document
        collections = []
        
        # Sélectionner les collections en fonction du type de document
        if document.metadata.get("document_type") == "presentation_ohada":
            collections.append("presentation_ohada")
            collections.append("plan_comptable")  # Ajouter aussi à la collection principale
        else:
            # Collections pour les documents standard
            collections.append("plan_comptable")  # Collection principale
            
            # Ajouter la collection de la partie
            if "partie" in document.metadata:
                collections.append(f"partie_{document.metadata['partie']}")
            
            # Ajouter les collections spécifiques basées sur le type
            document_type = document.metadata.get("document_type", "")
            if document_type == "chapitre":
                collections.append("chapitres")
        
        # Ajouter à chaque collection
        successful = True
        for collection in collections:
            try:
                self.vector_db.add_document(collection, document, embedding)
                print(f"Document ajouté à la collection '{collection}'")
            except Exception as e:
                print(f"Erreur lors de l'ajout à la collection '{collection}': {e}")
                successful = False
        
        # Mettre à jour la TOC avec des informations supplémentaires
        if successful and document.id:
            # Vérifier si l'entrée existe déjà dans la TOC
            if "lookup" not in self.toc:
                self.toc["lookup"] = {}
            
            if document.id not in self.toc["lookup"]:
                self.toc["lookup"][document.id] = {}
            
            # Mise à jour des informations
            self.toc["lookup"][document.id].update({
                "docx_path": docx_path,
                "title": document.metadata.get("title", ""),
                "document_type": document.metadata.get("document_type", ""),
                "partie": document.metadata.get("partie"),
                "chapitre": document.metadata.get("chapitre"),
                "id": document.id,
                "parent_id": document.metadata.get("parent_id")
            })
            
            # Sauvegarder la TOC mise à jour
            self.vector_db.save_toc(self.toc)
        
        return successful


def ingest_ohada_docx(docx_dir: str, vector_db: OhadaVectorDB, extensions: List[str] = ['.docx']):
    """Ingère les fichiers Word du répertoire dans la base vectorielle
    
    Args:
        docx_dir: Répertoire contenant les fichiers Word
        vector_db: Instance de la base de données vectorielle
        extensions: Extensions de fichiers à traiter
        
    Returns:
        Liste des fichiers traités avec succès
    """
    # Initialiser le processeur Word
    processor = OhadaWordProcessor(vector_db)
    
    # Rechercher les fichiers Word
    docx_files = []
    for ext in extensions:
        docx_files.extend(glob.glob(os.path.join(docx_dir, f"**/*{ext}"), recursive=True))
    
    # Filtrer pour ne garder que les chapitres (sans sections ni applications)
    chapitre_pattern = re.compile(r'chapitre_\d+\.docx$')
    chapitre_files = [f for f in docx_files if chapitre_pattern.search(os.path.basename(f))]
    
    # Trier les fichiers par partie et numéro de chapitre
    def extract_partie_chapitre(file_path):
        partie_match = re.search(r'partie_(\d+)', file_path)
        chapitre_match = re.search(r'chapitre_(\d+)', os.path.basename(file_path))
        partie = int(partie_match.group(1)) if partie_match else 0
        chapitre = int(chapitre_match.group(1)) if chapitre_match else 0
        return (partie, chapitre)
    
    chapitre_files.sort(key=extract_partie_chapitre)
    
    print(f"Nombre de fichiers de chapitres trouvés: {len(chapitre_files)}")
    
    # Ingérer chaque fichier
    successful_files = []
    failed_files = []
    
    for i, docx_path in enumerate(tqdm(chapitre_files, desc="Ingestion des chapitres")):
        print(f"\nIngestion {i+1}/{len(chapitre_files)}: {os.path.basename(docx_path)}")
        try:
            success = processor.ingest_docx_file(docx_path)
            if success:
                successful_files.append(docx_path)
            else:
                failed_files.append(docx_path)
        except Exception as e:
            print(f"Erreur lors de l'ingestion de {docx_path}: {e}")
            failed_files.append(docx_path)
    
    # Résumé
    print(f"\nIngestion terminée: {len(successful_files)} fichiers traités avec succès, {len(failed_files)} échecs")
    
    if failed_files:
        print("\nFichiers échoués:")
        for file in failed_files:
            print(f"  - {file}")
    
    # Afficher les statistiques de la base
    stats = vector_db.get_collection_stats()
    print("\nStatistiques des collections après ingestion:")
    for collection, data in stats.items():
        if "error" in data:
            print(f"  - {collection}: ERREUR - {data['error']}")
        else:
            print(f"  - {collection}: {data.get('count', 'N/A')} documents")
    
    return successful_files


def ingest_presentation_ohada(presentation_dir: str, vector_db: OhadaVectorDB):
    """Ingère les documents de présentation OHADA dans la base vectorielle
    
    Args:
        presentation_dir: Répertoire contenant les documents de présentation
        vector_db: Instance de la base de données vectorielle
        
    Returns:
        Liste des fichiers traités avec succès
    """
    # Initialiser le processeur Word
    processor = OhadaWordProcessor(vector_db)
    
    # Rechercher les fichiers Word dans le répertoire de présentation
    presentation_files = glob.glob(os.path.join(presentation_dir, "*.docx"))
    
    if not presentation_files:
        print(f"Aucun fichier de présentation trouvé dans {presentation_dir}")
        return []
    
    print(f"Nombre de fichiers de présentation trouvés: {len(presentation_files)}")
    
    # Ingérer chaque fichier
    successful_files = []
    failed_files = []
    
    for i, docx_path in enumerate(tqdm(presentation_files, desc="Ingestion des présentations OHADA")):
        print(f"\nIngestion {i+1}/{len(presentation_files)}: {os.path.basename(docx_path)}")
        try:
            success = processor.ingest_docx_file(docx_path)
            if success:
                successful_files.append(docx_path)
            else:
                failed_files.append(docx_path)
        except Exception as e:
            print(f"Erreur lors de l'ingestion de {docx_path}: {e}")
            failed_files.append(docx_path)
    
    # Résumé
    print(f"\nIngestion des présentations terminée: {len(successful_files)} fichiers traités avec succès, {len(failed_files)} échecs")
    
    if failed_files:
        print("\nFichiers échoués:")
        for file in failed_files:
            print(f"  - {file}")
    
    return successful_files


def build_toc_from_docx(docx_dir: str, toc_file: str):
    """Construit une table des matières à partir des fichiers Word
    
    Args:
        docx_dir: Répertoire contenant les fichiers Word
        toc_file: Chemin du fichier de sortie pour la table des matières
        
    Returns:
        Dictionnaire contenant la table des matières
    """
    # Structure de base
    toc = {
        "parties": [],
        "chapitres": [],
        "presentations": [],  # Nouvelle section pour les documents de présentation
        "lookup": {}
    }
    
    # Rechercher les répertoires de parties
    partie_dirs = [d for d in glob.glob(os.path.join(docx_dir, "partie_*")) if os.path.isdir(d)]
    
    for partie_dir in sorted(partie_dirs):
        partie_match = re.search(r'partie_(\d+)', partie_dir)
        if not partie_match:
            continue
        
        partie_num = int(partie_match.group(1))
        partie_id = f"partie_{partie_num}"
        
        # Ajouter la partie à la TOC
        partie_entry = {
            "type": "partie",
            "numero": partie_num,
            "titre": f"PARTIE {partie_num}",
            "id": partie_id
        }
        toc["parties"].append(partie_entry)
        toc["lookup"][partie_id] = partie_entry
        
        # Rechercher les fichiers de chapitres
        chapitre_files = glob.glob(os.path.join(partie_dir, "chapitre_*.docx"))
        for chapitre_file in sorted(chapitre_files):
            # Ignorer les fichiers de sections
            if "_section_" in chapitre_file or "_app_" in chapitre_file:
                continue
            
            chapitre_match = re.search(r'chapitre_(\d+)\.docx$', chapitre_file)
            if not chapitre_match:
                continue
            
            chapitre_num = int(chapitre_match.group(1))
            chapitre_id = f"{partie_id}_chapitre_{chapitre_num}"
            
            # Extraire le titre du chapitre du Word
            titre = f"CHAPITRE {chapitre_num}"
            try:
                doc = Document(chapitre_file)
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        title_match = re.search(r'(?:CHAPITRE|Chapitre)\s+\d+\s*[:\-]?\s*(.+)', paragraph.text)
                        if title_match:
                            titre = f"CHAPITRE {chapitre_num}: {title_match.group(1).strip()}"
                            break
            except Exception as e:
                print(f"Erreur lors de l'extraction du titre du chapitre {chapitre_file}: {e}")
            
            # Ajouter le chapitre à la TOC
            chapitre_entry = {
                "type": "chapitre",
                "numero": chapitre_num,
                "titre": titre,
                "id": chapitre_id,
                "parent_id": partie_id,
                "partie": partie_num,
                "docx_path": chapitre_file
            }
            toc["chapitres"].append(chapitre_entry)
            toc["lookup"][chapitre_id] = chapitre_entry
    
    # Rechercher les fichiers de présentation OHADA
    presentation_dir = os.path.join(os.path.dirname(docx_dir), "base_connaissances", "presentation_ohada")
    if os.path.exists(presentation_dir):
        presentation_files = glob.glob(os.path.join(presentation_dir, "*.docx"))
        
        for pres_file in sorted(presentation_files):
            file_name = os.path.basename(pres_file)
            file_base = os.path.splitext(file_name)[0]
            
            # Générer un ID pour le document de présentation
            clean_id = re.sub(r'[^a-z0-9_]', '_', file_base.lower())
            doc_id = f"presentation_ohada_{clean_id}"
            
            # Ajouter l'entrée de présentation à la TOC
            presentation_entry = {
                "type": "presentation_ohada",
                "titre": file_base,
                "id": doc_id,
                "docx_path": pres_file
            }
            toc["presentations"].append(presentation_entry)
            toc["lookup"][doc_id] = presentation_entry
    
    # Sauvegarder la TOC
    os.makedirs(os.path.dirname(toc_file), exist_ok=True)
    with open(toc_file, 'w', encoding='utf-8') as f:
        json.dump(toc, f, ensure_ascii=False, indent=2)
    
    print(f"Table des matières construite et sauvegardée dans {toc_file}")
    print(f"  - {len(toc['parties'])} parties")
    print(f"  - {len(toc['chapitres'])} chapitres")
    print(f"  - {len(toc['presentations'])} documents de présentation")
    
    return toc


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingesteur de Word pour le plan comptable OHADA")
    parser.add_argument("--docx-dir", default="./base_connaissances/plan_comptable/chapitres_word", 
                      help="Répertoire contenant les fichiers Word (défaut: ./base_connaissances/plan_comptable/chapitres_word)")
    parser.add_argument("--presentation-dir", default="./base_connaissances/presentation_ohada", 
                      help="Répertoire contenant les documents de présentation OHADA")
    parser.add_argument("--toc-file", default="./base_connaissances/plan_comptable/ohada_toc.json", 
                      help="Fichier de table des matières (défaut: ./base_connaissances/plan_comptable/ohada_toc.json)")
    parser.add_argument("--build-toc", action="store_true", 
                      help="Construire la table des matières")
    parser.add_argument("--reset", action="store_true", 
                      help="Réinitialiser la base de données vectorielle")
    parser.add_argument("--model", default=None, 
                      help="Modèle d'embedding à utiliser (si non spécifié, déterminé selon l'environnement)")
    parser.add_argument("--partie", type=int, choices=[1, 2, 3, 4], 
                      help="Traiter uniquement une partie spécifique (1-4)")
    parser.add_argument("--presentations-only", action="store_true",
                      help="Traiter uniquement les documents de présentation OHADA")
    parser.add_argument("--env", choices=["test", "production"], default=None,
                      help="Définir l'environnement (remplace la variable d'environnement OHADA_ENV)")
    
    args = parser.parse_args()
    
    # Définir l'environnement si spécifié
    if args.env:
        os.environ["OHADA_ENV"] = args.env
        print(f"Environnement défini: {args.env}")
    
    # Vérifier que les répertoires existent
    if not args.presentations_only and not os.path.exists(args.docx_dir):
        print(f"Erreur: Le répertoire {args.docx_dir} n'existe pas.")
        exit(1)
    
    if args.presentations_only and not os.path.exists(args.presentation_dir):
        print(f"Erreur: Le répertoire de présentation {args.presentation_dir} n'existe pas.")
        exit(1)
    
    # Construire la table des matières si demandé
    if args.build_toc:
        build_toc_from_docx(args.docx_dir, args.toc_file)
    
    # Initialiser la base vectorielle
    vector_db = OhadaVectorDB(toc_file=args.toc_file, embedding_model=args.model)
    
    # Réinitialiser la base si demandé
    if args.reset:
        response = input("Êtes-vous sûr de vouloir réinitialiser la base de données vectorielle? (y/n): ")
        if response.lower() == 'y':
            vector_db.reset_database()
    
    # Ingérer les documents de présentation OHADA si demandé
    if args.presentations_only or (not args.partie):
        if os.path.exists(args.presentation_dir):
            print(f"Traitement des documents de présentation OHADA dans {args.presentation_dir}")
            ingest_presentation_ohada(args.presentation_dir, vector_db)
        else:
            print(f"Avertissement: Le répertoire de présentation {args.presentation_dir} n'existe pas.")
    
    # Ingérer les fichiers Word si non exclusivement limité aux présentations
    if not args.presentations_only:
        if args.partie:
            # Filtrer uniquement la partie demandée
            partie_dir = os.path.join(args.docx_dir, f"partie_{args.partie}")
            if os.path.exists(partie_dir):
                print(f"Traitement uniquement de la partie {args.partie}")
                ingest_ohada_docx(partie_dir, vector_db)
            else:
                print(f"Erreur: Le répertoire de la partie {args.partie} n'existe pas: {partie_dir}")
        else:
            # Traiter toutes les parties
            ingest_ohada_docx(args.docx_dir, vector_db)