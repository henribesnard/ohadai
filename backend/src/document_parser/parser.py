"""
OHADA Document Parser - Extract text and metadata from Word documents
"""

import os
import hashlib
from pathlib import Path
from typing import Dict, Optional, List, Tuple
from datetime import datetime
import logging

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .extractor import HierarchyExtractor, HierarchyInfo

logger = logging.getLogger(__name__)


class OhadaDocumentParser:
    """
    Parse OHADA Word documents and extract:
    - Full text content
    - OHADA hierarchy (acte_uniforme, partie, chapitre, section, etc.)
    - Metadata (tags, references, dates)
    - Document type
    """

    def __init__(self):
        self.extractor = HierarchyExtractor()

    def parse_docx(self, file_path: str) -> Dict:
        """
        Parse a Word document (.docx) and extract all information

        Args:
            file_path: Path to .docx file

        Returns:
            Dictionary with extracted data:
            {
                'title': str,
                'content_text': str,
                'content_hash': str,
                'document_type': str,
                'collection': str | None,
                'sub_collection': str | None,
                'acte_uniforme': str | None,
                'livre': int | None,
                'titre': int | None,
                'partie': int | None,
                'chapitre': int | None,
                'section': int | None,
                'sous_section': str | None,
                'article': str | None,
                'alinea': int | None,
                'tags': List[str],
                'metadata': Dict,
                'date_publication': str | None,
                'page_count': int,
                'file_name': str,
                'file_size': int
            }

        Raises:
            FileNotFoundError: If file doesn't exist
            PackageNotFoundError: If file is not a valid .docx
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not file_path.suffix.lower() == '.docx':
            raise ValueError(f"File must be .docx format: {file_path}")

        logger.info(f"Parsing document: {file_path.name}")

        try:
            doc = Document(str(file_path))
        except PackageNotFoundError as e:
            raise PackageNotFoundError(f"Invalid .docx file: {file_path}") from e

        # Extract text content
        content_text = self._extract_text(doc)

        # Extract title (from first paragraph or filename)
        title = self._extract_title(doc, file_path)

        # Compute content hash for deduplication
        content_hash = self._compute_hash(content_text)

        # Extract hierarchy
        hierarchy = self.extractor.extract_hierarchy_from_text(content_text, title)

        # Extract document type
        document_type = self.extractor.extract_document_type(content_text, title)

        # Extract tags
        tags = self.extractor.extract_tags(content_text)

        # Extract references
        references = self.extractor.extract_references(content_text)

        # Extract date
        date_publication = self.extractor.extract_date_publication(content_text, title)

        # Extract collection and sub_collection from file path
        collection, sub_collection = self._extract_collection_from_path(file_path)

        # Get file metadata
        file_stats = file_path.stat()

        # Build metadata
        metadata = {
            'file_name': file_path.name,
            'file_size': file_stats.st_size,
            'file_modified': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            'file_path': str(file_path),
            'collection': collection,
            'sub_collection': sub_collection,
            'references': references,
            'parsed_at': datetime.now().isoformat(),
            'parser_version': '1.1'
        }

        result = {
            'title': title,
            'content_text': content_text,
            'content_hash': content_hash,
            'document_type': document_type,
            'collection': collection,
            'sub_collection': sub_collection,
            'acte_uniforme': hierarchy.acte_uniforme,
            'livre': hierarchy.livre,
            'titre': hierarchy.titre,
            'partie': hierarchy.partie,
            'chapitre': hierarchy.chapitre,
            'section': hierarchy.section,
            'sous_section': hierarchy.sous_section,
            'article': hierarchy.article,
            'alinea': hierarchy.alinea,
            'tags': tags,
            'metadata': metadata,
            'date_publication': date_publication,
            'page_count': self._estimate_page_count(content_text),
            'file_name': file_path.name,
            'file_size': file_stats.st_size
        }

        logger.info(f"Parsed: {title} - Type: {document_type} - Hierarchy: "
                   f"Partie {hierarchy.partie}, Chapitre {hierarchy.chapitre}")

        return result

    def _extract_text(self, doc: Document) -> str:
        """
        Extract all text from document paragraphs and tables

        Args:
            doc: python-docx Document object

        Returns:
            Full text content
        """
        paragraphs_text = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs_text.append(text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs_text.append(text)

        full_text = '\n'.join(paragraphs_text)
        return full_text

    def _extract_title(self, doc: Document, file_path: Path) -> str:
        """
        Extract document title from first paragraph or filename

        Priority:
        1. First paragraph if it looks like a title (< 200 chars, all caps or title case)
        2. Filename without extension

        Args:
            doc: python-docx Document object
            file_path: Path to file

        Returns:
            Document title
        """
        # Try first paragraph
        if doc.paragraphs and len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0].text.strip()

            # Check if it looks like a title
            if first_para and len(first_para) < 200:
                # Check for common title patterns
                if (first_para.isupper() or
                    first_para.istitle() or
                    any(word in first_para.lower() for word in
                        ['acte uniforme', 'chapitre', 'partie', 'syscohada'])):
                    return first_para

        # Fallback to filename
        title = file_path.stem.replace('_', ' ').replace('-', ' ')
        return title.title()

    def _compute_hash(self, content: str) -> str:
        """
        Compute SHA-256 hash of content for deduplication

        Args:
            content: Text content

        Returns:
            Hex digest of SHA-256 hash
        """
        return hashlib.sha256(content.encode('utf-8')).hexdigest()

    def _estimate_page_count(self, text: str) -> int:
        """
        Estimate page count based on text length

        Assumes ~500 words per page, ~5 chars per word

        Args:
            text: Full text content

        Returns:
            Estimated page count
        """
        chars_per_page = 2500  # 500 words * 5 chars
        page_count = max(1, len(text) // chars_per_page)
        return page_count

    def _extract_collection_from_path(self, file_path: Path) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract collection and sub_collection from file path

        Analyzes the directory structure to determine:
        - collection: Main category (actes_uniformes, plan_comptable, presentation_ohada, etc.)
        - sub_collection: Subcategory (specific acte name, partie number, etc.)

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (collection, sub_collection)

        Examples:
            base_connaissances/actes_uniformes/droit commercial général/Livre_1.docx
            → ('Actes Uniformes', 'Droit Commercial Général')

            base_connaissances/plan_comptable/chapitres_word/partie_1/chapitre_1.docx
            → ('Plan Comptable SYSCOHADA', 'Partie 1')

            base_connaissances/presentation_ohada/Introduction.docx
            → ('Présentation OHADA', None)
        """
        parts = file_path.parts

        # Try to find base_connaissances in path
        try:
            base_idx = parts.index('base_connaissances')
        except ValueError:
            # Not in base_connaissances structure, return None
            logger.warning(f"File not in base_connaissances structure: {file_path}")
            return None, None

        # Extract path after base_connaissances
        remaining_parts = parts[base_idx + 1:]

        if not remaining_parts:
            return None, None

        # First level after base_connaissances is the collection
        collection_dir = remaining_parts[0]

        # Map directory names to collection names
        collection_mapping = {
            'actes_uniformes': 'Actes Uniformes',
            'plan_comptable': 'Plan Comptable SYSCOHADA',
            'presentation_ohada': 'Présentation OHADA',
            'jurisprudence': 'Jurisprudence',
            'doctrine': 'Doctrine',
            'reglements': 'Règlements',
            # Ajoutez vos collections personnalisées ici:
            # 'guides_pratiques': 'Guides Pratiques et Outils',
            # 'circulaires_ohada': 'Circulaires OHADA',
            # 'notes_explicatives': 'Notes Explicatives et Commentaires',
        }

        collection = collection_mapping.get(collection_dir, collection_dir.replace('_', ' ').title())

        # Second level is the sub_collection
        sub_collection = None

        if len(remaining_parts) > 1:
            # For actes_uniformes: sub_collection is the acte name
            if collection_dir == 'actes_uniformes':
                sub_collection = remaining_parts[1].title()

            # For plan_comptable: extract partie number
            elif collection_dir == 'plan_comptable':
                # Path might be: plan_comptable/chapitres_word/partie_1/...
                for part in remaining_parts[1:]:
                    if part.startswith('partie_'):
                        partie_num = part.replace('partie_', '')
                        sub_collection = f"Partie {partie_num}"
                        break

            # For other collections: use second level directory
            else:
                sub_collection = remaining_parts[1].replace('_', ' ').title()

        logger.debug(f"Extracted: collection='{collection}', sub_collection='{sub_collection}'")

        return collection, sub_collection

    def parse_directory(self, directory_path: str, pattern: str = "*.docx") -> List[Dict]:
        """
        Parse all .docx files in a directory

        Args:
            directory_path: Path to directory containing .docx files
            pattern: Glob pattern for files (default: "*.docx")

        Returns:
            List of parsed document dictionaries

        Raises:
            FileNotFoundError: If directory doesn't exist
        """
        directory = Path(directory_path)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        if not directory.is_dir():
            raise ValueError(f"Not a directory: {directory}")

        logger.info(f"Parsing directory: {directory}")

        documents = []
        files = list(directory.glob(pattern))

        logger.info(f"Found {len(files)} files matching {pattern}")

        for file_path in files:
            try:
                doc_data = self.parse_docx(str(file_path))
                documents.append(doc_data)
            except Exception as e:
                logger.error(f"Failed to parse {file_path.name}: {e}")
                continue

        logger.info(f"Successfully parsed {len(documents)}/{len(files)} documents")

        return documents

    def validate_document_data(self, doc_data: Dict) -> List[str]:
        """
        Validate parsed document data

        Args:
            doc_data: Parsed document dictionary

        Returns:
            List of validation warnings (empty if valid)
        """
        warnings = []

        # Required fields
        if not doc_data.get('title'):
            warnings.append("Missing title")

        if not doc_data.get('content_text'):
            warnings.append("Missing content_text")

        if len(doc_data.get('content_text', '')) < 100:
            warnings.append("Content too short (< 100 chars)")

        # Check hierarchy consistency
        if doc_data.get('section') and not doc_data.get('chapitre'):
            warnings.append("Section without chapitre")

        if doc_data.get('article') and not doc_data.get('chapitre'):
            warnings.append("Article without chapitre")

        # Document type validation
        if doc_data.get('document_type') == 'other':
            warnings.append("Could not determine specific document type")

        return warnings
